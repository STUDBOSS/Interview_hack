# # # main.py

# # from audio.recorder import record_audio
# # from transcriber.whisper_transcriber import transcribe_audio
# # from gpt_responder.gpt_responder import ask_gpt


# # import time

# # def main():
# #     print("🧠 Interview Assistant started. Listening for questions... (Press Ctrl+C to stop)")

# #     try:
# #         while True:
# #             # Step 1: Record a chunk of audio
# #             audio = record_audio(duration=5)

# #             # Step 2: Transcribe it to text
# #             text = transcribe_audio(audio)

# #             if text:
# #                 print(f"\n🗣️ Question Heard: {text}")

# #                 # Step 3: Send to GPT
# #                 answer = ask_gpt(text)

# #                 # Step 4: Display the response
# #                 print(f"🤖 GPT Suggests: {answer}\n")

# #             else:
# #                 print("🔇 Silence detected. Waiting...\n")

# #             time.sleep(1)  # Small delay to prevent overlapping recordings

# #     except KeyboardInterrupt:
# #         print("\n🛑 Stopped by user. Exiting...")

# # if __name__ == "__main__":
# #     main()


# # import sounddevice as sd
# # import soundfile as sf
# # import tempfile
# # import os
# # from transcriber.whisper_transcriber import transcribe_audio
# # import openai

# # # Initialize OpenAI client with your API Key
# # client = openai.OpenAI(
# #     api_key="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"
# # )

# # def record_audio(duration=5, samplerate=16000):
# #     print("Recording...")
# #     recording = sd.rec(int(duration * samplerate), samplerate=samplerate, channels=1)
# #     sd.wait()
# #     print("Recording complete.")
# #     tmpfile = tempfile.NamedTemporaryFile(delete=False, suffix=".wav")
# #     sf.write(tmpfile.name, recording, samplerate)
# #     return tmpfile.name

# # def ask_gpt(question):
# #     try:
# #         response = client.chat.completions.create(
# #             model="gpt-3.5-turbo",
# #             messages=[
# #                 {"role": "system", "content": "You are a helpful AI Interview Assistant."},
# #                 {"role": "user", "content": question}
# #             ]
# #         )
# #         answer = response.choices[0].message.content.strip()
# #         return answer
# #     except Exception as e:
# #         print(f"[ERROR] GPT request failed: {e}")
# #         return None

# # def main():
# #     while True:
# #         print("\n🧠 Interview Assistant started. Listening for questions... (Press Ctrl+C to stop)")
# #         try:
# #             audio = record_audio(duration=5)
# #             text = transcribe_audio(audio)
# #             if text.strip() == "":
# #                 print("🔇 Silence detected. Waiting...")
# #                 continue

# #             print(f"\n🗣️ Question Heard: {text}")

# #             answer = ask_gpt(text)
# #             if answer:
# #                 print(f"\n🤖 GPT Suggests: {answer}")
# #             else:
# #                 print("[ERROR] Failed to get response from GPT.")

# #             os.remove(audio)

# #         except KeyboardInterrupt:
# #             print("\n🛑 Stopped by user. Exiting...")
# #             break
# #         except Exception as e:
# #             print(f"[ERROR] {e}")

# # if __name__ == "__main__":
# #     main()


# import requests
# from transcriber import whisper_transcriber

# def send_message_to_gpt(message):
#     url = "https://chatgpt-42.p.rapidapi.com/chat"
#     payload = {
#         "messages": [{"role": "user", "content": message}],
#         "model": "gpt-4o-mini"
#     }
#     headers = {
#         "x-rapidapi-host": "chatgpt-42.p.rapidapi.com",
#         "x-rapidapi-key": "sk-efgh5678abcd1234efgh5678abcd1234efgh5678",  # 👈 ADD your API KEY here!
#         "Content-Type": "application/json"
#     }
#     response = requests.post(url, json=payload, headers=headers)

#     if response.status_code == 200:
#         return response.json()
#     else:
#         print(f"Error: {response.status_code}, {response.text}")
#         return None

# def main():
#     print("🧠 Interview Assistant started. Listening for questions... (Press Ctrl+C to stop)")

#     while True:
#         try:
#             audio = whisper_transcriber.record_audio()
#             text = whisper_transcriber.transcribe_audio(audio)
            
#             if text:
#                 print(f"🗣️ Question Heard: {text}")
#                 gpt_response = send_message_to_gpt(text)
                
#                 if gpt_response:
#                     print("🤖 GPT Suggests:", gpt_response['choices'][0]['message']['content'])
#             else:
#                 print("🔇 Silence detected. Waiting...")

#         except KeyboardInterrupt:
#             print("\n🛑 Stopped by user. Exiting...")
#             break

# if __name__ == "__main__":
#     main()


# import time
# from transcriber import whisper_transcriber
# import requests

# RAPIDAPI_KEY = "2da17ffaa1mshc9f9e89ffaad466p1514d5jsn230922b2d57c"
# RAPIDAPI_URL = "https://chatgpt-42.p.rapidapi.com/chat"

# def ask_gpt(prompt):
#     url = RAPIDAPI_URL

#     payload = {
#         "messages": [
#             {
#                 "role": "user",
#                 "content": prompt
#             }
#         ],
#         "model": "gpt-4o-mini"
#     }
#     headers = {
#         "x-rapidapi-key": RAPIDAPI_KEY,
#         "x-rapidapi-host": "chatgpt-42.p.rapidapi.com",
#         "Content-Type": "application/json"
#     }

#     try:
#         response = requests.post(url, json=payload, headers=headers)
#         response.raise_for_status()
#         data = response.json()
#         return data.get('result') or data
#     except Exception as e:
#         print(f"❌ Error communicating with GPT: {e}")
#         return None

# def main():
#     print("🧠 Interview Assistant started. Listening for questions... (Press Ctrl+C to stop)")

#     try:
#         while True:
#             audio_path = whisper_transcriber.record_audio(duration=5)
#             question = whisper_transcriber.transcribe_audio(audio_path)

#             if not question:
#                 print("🔇 Silence or unclear speech detected. Waiting...")
#                 time.sleep(1)
#                 continue

#             print(f"\n🗣️ Question Heard: {question}")

#             response = ask_gpt(question)

#             if response:
#                 print(f"\n🤖 GPT Suggests: {response}")
#             else:
#                 print("⚠️ No response received from GPT.")

#             print("\nRecording next question...\n")
#             time.sleep(1)

#     except KeyboardInterrupt:
#         print("\n🛑 Stopped by user. Exiting...")

# if __name__ == "__main__":
#     main()


# import time
# from transcriber import whisper_transcriber
# from resume_processor import ResumeAssistant  # 👈 new import
# import requests

# def main():
#     print("🧠 Interview Assistant started. Loading resume...")

#     # Initialize the assistant with resume
#     resume_path = r"C:\VS Code\Interview\interview-hack-gpt\Mohit-Paradkar-B.E.-ArtificialIntelligence&MachineLearning-2025-03-16-06-58-00-916258.pdf"  # 👈 update with your actual file name or path
#     assistant = ResumeAssistant(resume_path)

#     print("📄 Resume loaded. Listening for questions... (Press Ctrl+C to stop)")

#     try:
#         while True:
#             audio_path = whisper_transcriber.record_audio(duration=5)
#             question = whisper_transcriber.transcribe_audio(audio_path)

#             if not question:
#                 print("🔇 Silence or unclear speech detected. Waiting...")
#                 time.sleep(1)
#                 continue

#             print(f"\n🗣️ Question Heard: {question}")

#             # Ask GPT using the ResumeAssistant
#             response = assistant.ask_question(question)

#             if response:
#                 print(f"\n🤖 GPT Suggests: {response}")
#             else:
#                 print("⚠️ No response received from GPT.")

#             print("\nRecording next question...\n")
#             time.sleep(1)

#     except KeyboardInterrupt:
#         print("\n🛑 Stopped by user. Exiting...")

# if __name__ == "__main__":
#     main()


import time
from caption_reader import live_caption_listener  # <-- NEW
import requests

# RAPIDAPI_KEY = "2da17ffaa1mshc9f9e89ffaad466p1514d5jsn230922b2d57c"
# RAPIDAPI_URL = "https://chatgpt-42.p.rapidapi.com/chat"

# def ask_gpt(prompt):
#     url = RAPIDAPI_URL
#     payload = {
#         "messages": [
#             {
#                 "role": "user",
#                 "content": prompt
#             }
#         ],
#         "model": "gpt-4o-mini"
#     }
#     headers = {
#         "x-rapidapi-key": RAPIDAPI_KEY,
#         "x-rapidapi-host": "chatgpt-42.p.rapidapi.com",
#         "Content-Type": "application/json"
#     }

#     try:
#         response = requests.post(url, json=payload, headers=headers)
#         response.raise_for_status()
#         data = response.json()
#         return data.get('result') or data
#     except Exception as e:
#         print(f"❌ Error communicating with GPT: {e}")
#         return None

# def handle_caption(caption_text):
#     print(f"\n🗣️ Question Heard: {caption_text}")
#     response = ask_gpt(caption_text)
#     if response:
#         print(f"\n🤖 GPT Suggests: {response}")
#     else:
#         print("⚠️ No response received from GPT.")
#     print("\nListening for next question...\n")

# def main():
#     print("🧠 Interview Assistant started with Live Captions.")
#     try:
#         live_caption_listener(handle_caption)  # <-- New loop
#     except KeyboardInterrupt:
#         print("\n🛑 Stopped by user. Exiting...")

# if __name__ == "__main__":
#     main()


# import requests

# # API details
# RAPIDAPI_KEY = "2da17ffaa1mshc9f9e89ffaad466p1514d5jsn230922b2d57c"
# RAPIDAPI_URL = "https://chatgpt-42.p.rapidapi.com/chat"

# def ask_gpt(prompt):
#     """Ask GPT for a response using the provided prompt."""
#     url = RAPIDAPI_URL
#     payload = {
#         "messages": [
#             {"role": "user", "content": prompt}
#         ],
#         "model": "gpt-4o-mini"
#     }
#     headers = {
#         "x-rapidapi-key": RAPIDAPI_KEY,
#         "x-rapidapi-host": "chatgpt-42.p.rapidapi.com",
#         "Content-Type": "application/json"
#     }

#     try:
#         response = requests.post(url, json=payload, headers=headers)
#         response.raise_for_status()  # Ensure we catch HTTP errors
#         data = response.json()
#         return data.get('result') or data
#     except Exception as e:
#         print(f"❌ Error communicating with GPT: {e}")
#         return None

# def handle_caption(caption_text, resume_text):
#     """Handle the captured caption by sending it to GPT."""
#     print(f"\n🗣️ Question Heard: {caption_text}")
    
#     # Combine resume context with the user's question
#     prompt = f"You are preparing interview answers based on the following resume:\n{resume_text}\nQuestion: {caption_text}\nAnswer:"
    
#     # Ask GPT for a response
#     response = ask_gpt(prompt)
    
#     if response:
#         print(f"\n🤖 GPT Suggests: {response}")
#     else:
#         print("⚠️ No response received from GPT.")
#     print("\nListening for next question...\n")
# # Function to extract resume text (from PDF or DOCX)
# def extract_text_from_pdf(pdf_path):
#     import PyPDF2
#     with open(pdf_path, "rb") as file:
#         reader = PyPDF2.PdfReader(file)
#         text = ""
#         for page in reader.pages:
#             text += page.extract_text()
#     return text

# def extract_text_from_docx(docx_path):
#     import docx
#     doc = docx.Document(docx_path)
#     text = ""
#     for para in doc.paragraphs:
#         text += para.text + "\n"
#     return text

# # Example workflow for extracting the resume text and processing captions
# def main(resume_path):
#     # Extract resume text based on file format
#     if resume_path.endswith(".pdf"):
#         resume_text = extract_text_from_pdf(resume_path)
#     elif resume_path.endswith(".docx"):
#         resume_text = extract_text_from_docx(resume_path)
#     else:
#         print("Unsupported file format.")
#         return
    
#     print("Resume successfully loaded. Preparing to listen for questions...\n")
    
#     # Simulating the caption capture process (you would replace this with real-time speech capture)
#     while True:
#         caption = input("Enter a question for GPT (or type 'exit' to quit): ")
        
#         if caption.lower() == 'exit':
#             break
        
#         handle_caption(caption, resume_text)

# # Example usage
# main("Mohit-Paradkar-B.E.-ArtificialIntelligence&MachineLearning-2025-03-16-06-58-00-916258.pdf")  # Replace with your resume file path


import requests
import time

# API details
RAPIDAPI_KEY = "2da17ffaa1mshc9f9e89ffaad466p1514d5jsn230922b2d57c"
RAPIDAPI_URL = "https://chatgpt-42.p.rapidapi.com/chat"

def ask_gpt(prompt):
    """Ask GPT for a response using the provided prompt."""
    url = RAPIDAPI_URL
    payload = {
        "messages": [
            {"role": "user", "content": prompt}
        ],
        "model": "gpt-4o-mini"
    }
    headers = {
        "x-rapidapi-key": RAPIDAPI_KEY,
        "x-rapidapi-host": "chatgpt-42.p.rapidapi.com",
        "Content-Type": "application/json"
    }

    try:
        print(f"Sending prompt to GPT: {prompt[:100]}...")  # Debugging: print the first 100 characters of the prompt
        response = requests.post(url, json=payload, headers=headers)
        response.raise_for_status()  # Ensure we catch HTTP errors
        data = response.json()
        result = data.get('result', None)
        if result:
            return result
        else:
            return "No valid result returned."
    except Exception as e:
        print(f"❌ Error communicating with GPT: {e}")
        return None

def handle_caption(caption_text, resume_text):
    """Handle the captured caption by sending it to GPT."""
    print(f"\n🗣️ Question Heard: {caption_text}")
    
    # Construct the prompt with the resume context and the question
    prompt = f"You are preparing interview answers based on the following resume:\n{resume_text}\nQuestion: {caption_text}\nAnswer:"
    
    # Ask GPT for a response
    response = ask_gpt(prompt)
    
    if response:
        print(f"\n🤖 GPT Suggests: {response}")
    else:
        print("⚠️ No response received from GPT.")
    print("\nListening for next question...\n")

# Function to extract resume text (from PDF or DOCX)
def extract_text_from_pdf(pdf_path):
    import PyPDF2
    with open(pdf_path, "rb") as file:
        reader = PyPDF2.PdfReader(file)
        text = ""
        for page in reader.pages:
            text += page.extract_text()
    return text

def extract_text_from_docx(docx_path):
    import docx
    doc = docx.Document(docx_path)
    text = ""
    for para in doc.paragraphs:
        text += para.text + "\n"
    return text

# Example workflow for extracting the resume text and processing captions
def main(resume_path):
    # Extract resume text based on file format
    if resume_path.endswith(".pdf"):
        resume_text = extract_text_from_pdf(resume_path)
    elif resume_path.endswith(".docx"):
        resume_text = extract_text_from_docx(resume_path)
    else:
        print("Unsupported file format.")
        return
    
    print("Resume successfully loaded. Preparing to listen for questions...\n")
    
    # Simulating the caption capture process (you would replace this with real-time speech capture)
    while True:
        caption = input("Enter a question for GPT (or type 'exit' to quit): ")
        
        if caption.lower() == 'exit':
            break
        
        handle_caption(caption, resume_text)

# Example usage
main("Mohit-Paradkar-B.E.-ArtificialIntelligence&MachineLearning-2025-03-16-06-58-00-916258.pdf")  # Replace with your resume file path
